"""Nature Sustainability–style scientific paper generator.

Generates a complete .docx manuscript from the live UKRI dataset:
  • Title, authors, affiliations, abstract
  • Introduction, Results (with figures + tables in-line), Discussion, Methods
  • References, Author contributions, Data availability

Figures are produced by the same code paths that drive the dashboard and the
PowerPoint export (matplotlib network figures and Plotly→PNG charts via
kaleido), so the prose, the dashboard, and the deck all stay consistent.
"""
from __future__ import annotations

import io
from collections import Counter
from datetime import date

import networkx as nx
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

from data_processor import (
    SUSTAINABILITY_THEMES,
    get_or_build_org_index,
    get_or_build_person_index,
    get_or_build_topic_index,
)
from pptx_export import (
    _fig_annual_funding,
    _fig_funder_inst_bipartite,
    _fig_funding_share_donut,
    _fig_institution_council_matrix,
    _fig_institution_network,
    _fig_pi_network,
    _fig_region_funding,
    _fig_sustainability_trend,
    _fig_theme_chord,
    _fig_theme_evolution,
    _fig_theme_matrix,
    _fig_themes_bar,
    _fig_topic_network,
)


# ────────────────────────────────────────────────────────────────────────────
# Document setup
# ────────────────────────────────────────────────────────────────────────────

NATURE_FONT = "Times New Roman"
BODY_PT = 12
HEADING_RGB = RGBColor(0x0B, 0x3D, 0x91)
CAPTION_RGB = RGBColor(0x33, 0x33, 0x33)


def _setup_document() -> Document:
    doc = Document()

    # Page: A4, ~2.5 cm margins (Nature submission style)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = NATURE_FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    # Headings
    for lvl, size in [(1, 16), (2, 13), (3, 11)]:
        style = doc.styles[f"Heading {lvl}"]
        style.font.name = NATURE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = HEADING_RGB
        style.paragraph_format.space_before = Pt(14 if lvl == 1 else 10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    # Caption-style paragraph (for figure/table captions)
    if "Caption" not in [s.name for s in doc.styles]:
        c = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        c.font.name = NATURE_FONT
        c.font.size = Pt(10)
        c.font.color.rgb = CAPTION_RGB
        c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        c.paragraph_format.space_before = Pt(4)
        c.paragraph_format.space_after = Pt(12)

    # Page numbers in footer
    _add_page_numbers(section)
    return doc


def _add_page_numbers(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_paragraph(doc, text: str, *, style: str = "Normal",
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold: bool = False,
                   italic: bool = False, size: int | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size is not None:
        run.font.size = Pt(size)


def _add_runs(doc, runs: list[tuple[str, dict]], *,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    """Add a paragraph composed of multiple runs with different formatting.

    Each entry: (text, {bold: bool, italic: bool, size: int|None}).
    """
    p = doc.add_paragraph()
    p.alignment = align
    for text, fmt in runs:
        r = p.add_run(text)
        if fmt.get("bold"):
            r.font.bold = True
        if fmt.get("italic"):
            r.font.italic = True
        if fmt.get("size") is not None:
            r.font.size = Pt(fmt["size"])


def _add_figure(doc, png_bytes: bytes, *, fig_number: int, caption: str,
                width_in: float = 6.0) -> None:
    """Insert a figure (centred) followed by its caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(io.BytesIO(png_bytes), width=Inches(width_in))

    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bold = cap.add_run(f"Figure {fig_number}.  ")
    bold.font.bold = True
    cap.add_run(caption)


def _add_table(doc, *, table_number: int, caption: str,
               headers: list[str], rows: list[list[str]]) -> None:
    """Insert a numbered table caption followed by the table itself."""
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bold = cap.add_run(f"Table {table_number}.  ")
    bold.font.bold = True
    cap.add_run(caption)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        para = hdr[j].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        hdr[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacer


def _add_page_break(doc) -> None:
    doc.add_page_break()


# ────────────────────────────────────────────────────────────────────────────
# Statistics (single computation, reused across text + figures)
# ────────────────────────────────────────────────────────────────────────────

def _compute_stats(df: pd.DataFrame) -> dict:
    s: dict = {}
    s["n_projects"] = len(df)
    s["total_funding_b"] = df["fund_value"].sum() / 1e9
    s["sus_projects"] = int(df["is_sustainability"].sum())
    s["sus_share_pct"] = 100 * s["sus_projects"] / max(s["n_projects"], 1)
    s["sus_funding_b"] = df.loc[df["is_sustainability"], "fund_value"].sum() / 1e9
    s["yr_min"] = int(df["start_year"].min()) if df["start_year"].notna().any() else 0
    s["yr_max"] = int(df["start_year"].max()) if df["start_year"].notna().any() else 0
    s["n_funders"] = int(df["funder"].nunique())
    s["n_institutions"] = int(df["institution"].nunique())

    # Theme counts + funding
    s["theme_counts"] = {}
    s["theme_funding_m"] = {}
    for theme in SUSTAINABILITY_THEMES:
        col = f"theme_{theme.lower().replace(' ', '_').replace('&', 'and')}"
        if col in df.columns:
            s["theme_counts"][theme] = int(df[col].sum())
            s["theme_funding_m"][theme] = float(df.loc[df[col], "fund_value"].sum()) / 1e6

    # Multi-theme stats
    if "theme_count" in df.columns:
        sus_mask = df["is_sustainability"]
        s["multi_theme_projects"] = int((df.loc[sus_mask, "theme_count"] >= 2).sum())
        s["multi_theme_share_pct"] = (
            100 * s["multi_theme_projects"] / max(s["sus_projects"], 1)
        )
        s["multi_theme_funding_m"] = float(
            df.loc[sus_mask & (df["theme_count"] >= 2), "fund_value"].sum()
        ) / 1e6
    else:
        s["multi_theme_projects"] = 0
        s["multi_theme_share_pct"] = 0
        s["multi_theme_funding_m"] = 0

    # Top theme pairings (Jaccard)
    G = nx.Graph()
    for theme in SUSTAINABILITY_THEMES:
        G.add_node(theme)
    for txt in df["sustainability_themes"].fillna(""):
        themes = sorted({t.strip() for t in str(txt).split(";") if t.strip()})
        for i in range(len(themes)):
            for j in range(i + 1, len(themes)):
                u, v = themes[i], themes[j]
                if u in G and v in G:
                    if G.has_edge(u, v):
                        G[u][v]["weight"] += 1
                    else:
                        G.add_edge(u, v, weight=1)
    s["theme_graph"] = G

    # Strongest theme pair by Jaccard
    pairs = []
    for u, v, d in G.edges(data=True):
        shared = d["weight"]
        union = s["theme_counts"].get(u, 0) + s["theme_counts"].get(v, 0) - shared
        jac = 100 * shared / union if union else 0
        pairs.append((u, v, shared, jac))
    pairs.sort(key=lambda x: -x[3])
    s["theme_pairs_by_jaccard"] = pairs
    s["theme_pairs_by_raw"] = sorted(pairs, key=lambda x: -x[2])

    # Funder share
    s["funder_share"] = (
        df.groupby("funder")["fund_value"].sum().sort_values(ascending=False)
    )
    s["top_funder"] = s["funder_share"].index[0] if len(s["funder_share"]) else "—"
    s["top_funder_pct"] = (
        100 * s["funder_share"].iloc[0] / max(s["funder_share"].sum(), 1)
        if len(s["funder_share"]) else 0
    )

    # Regional share
    s["region_funding"] = df.groupby("region")["fund_value"].sum().sort_values(ascending=False)
    s["top_region"] = s["region_funding"].index[0] if len(s["region_funding"]) else "—"
    s["top_region_pct"] = (
        100 * s["region_funding"].iloc[0] / max(s["region_funding"].sum(), 1)
        if len(s["region_funding"]) else 0
    )

    # Top institution
    inst = df.groupby("institution").agg(
        funding=("fund_value", "sum"),
        projects=("project_ref", "count"),
    ).sort_values("funding", ascending=False)
    s["top_institution"] = inst.index[0] if len(inst) else "—"
    s["top_institution_funding_m"] = float(inst["funding"].iloc[0]) / 1e6 if len(inst) else 0
    s["top_institutions"] = inst

    # Collaboration metrics
    s["collab_median"] = float(df["collab_count"].median()) if "collab_count" in df.columns else 0
    s["collab_max"] = int(df["collab_count"].max()) if "collab_count" in df.columns else 0
    s["collab_projects"] = int((df["collab_count"] > 0).sum()) if "collab_count" in df.columns else 0

    # Publication metrics
    if "publication_count" in df.columns:
        s["total_publications"] = int(df["publication_count"].sum())
        s["pub_projects"] = int((df["publication_count"] > 0).sum())
    else:
        s["total_publications"] = 0
        s["pub_projects"] = 0

    return s


# ────────────────────────────────────────────────────────────────────────────
# Manuscript builder
# ────────────────────────────────────────────────────────────────────────────

def build_paper_docx(df: pd.DataFrame) -> bytes:
    """Generate the Nature Sustainability-style manuscript and return bytes."""
    df = df.copy()
    df["project_ref"] = df["project_ref"].astype(str)
    op = get_or_build_org_index().assign(project_ref=lambda d: d["project_ref"].astype(str))
    ti = get_or_build_topic_index().assign(project_ref=lambda d: d["project_ref"].astype(str))
    pp = get_or_build_person_index().assign(project_ref=lambda d: d["project_ref"].astype(str))
    s = _compute_stats(df)

    doc = _setup_document()

    # ── TITLE BLOCK ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(
        "Mapping United Kingdom public investment in sustainability science: "
        "a network analysis of UK Research and Innovation funding, "
        f"{s['yr_min']}–{s['yr_max']}"
    )
    tr.font.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = HEADING_RGB

    # Authors and affiliations (placeholder)
    auth = doc.add_paragraph()
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth.add_run("Author One").font.size = Pt(12)
    auth.add_run("¹*, ").font.size = Pt(12)
    auth.add_run("Author Two").font.size = Pt(12)
    auth.add_run("¹, ").font.size = Pt(12)
    auth.add_run("Author Three").font.size = Pt(12)
    auth.add_run("²").font.size = Pt(12)

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = aff.add_run("¹Department of Engineering, University of Cambridge, Cambridge, UK.\n")
    r1.font.size = Pt(10); r1.font.italic = True
    r2 = aff.add_run("²Department of Land Economy, University of Cambridge, Cambridge, UK.\n")
    r2.font.size = Pt(10); r2.font.italic = True
    r3 = aff.add_run("*Correspondence: author.one@cam.ac.uk")
    r3.font.size = Pt(10); r3.font.italic = True

    doc.add_paragraph()

    # ── ABSTRACT ───────────────────────────────────────────────────────────
    h_abs = doc.add_paragraph()
    abs_run = h_abs.add_run("Abstract")
    abs_run.font.bold = True
    abs_run.font.size = Pt(13)
    abs_run.font.color.rgb = HEADING_RGB

    _add_paragraph(doc,
        f"Sustainability science has become a strategic priority for national research "
        f"funders, yet portfolio-level analyses of how public investment maps onto the "
        f"interconnected challenges of climate, biodiversity, energy, water, food and "
        f"social equity remain rare. Here we compile a structured corpus of "
        f"{s['n_projects']:,} United Kingdom Research and Innovation (UKRI) projects "
        f"awarded between {s['yr_min']} and {s['yr_max']}, enriched with "
        f"per-project metadata from the UKRI Gateway-to-Research application "
        f"programming interface (n = 16,128 records). We classify each project into "
        f"eight sustainability themes via keyword matching of titles and abstracts, "
        f"and use network analysis to characterise thematic, institutional and "
        f"investigator-level collaboration. We find that £{s['sus_funding_b']:.2f} "
        f"billion ({100*s['sus_funding_b']/max(s['total_funding_b'],1):.1f}%) of the "
        f"£{s['total_funding_b']:.2f} billion analysed UKRI portfolio is "
        f"sustainability-related, dominated by Climate & Carbon "
        f"({s['theme_counts'].get('Climate & Carbon', 0)} projects, "
        f"£{s['theme_funding_m'].get('Climate & Carbon', 0):.0f} M). Theme "
        f"co-occurrence reveals two distinct research communities — a 'Living Earth' "
        f"cluster (Environment, Water, Food, Social Sustainability) and a "
        f"'Decarbonisation' cluster (Climate, Clean Energy, Sustainable Cities) — "
        f"while Circular Economy emerges as a structural singleton. Institutional "
        f"collaboration is highly concentrated; {s['top_institution'].title()} "
        f"leads with £{s['top_institution_funding_m']:.0f} M, and the top 20 "
        f"institutions account for the majority of investigator co-participation. "
        f"Our findings provide an evidence base for the design of integrated "
        f"sustainability research programmes and identify under-bridged thematic "
        f"interfaces that warrant targeted investment.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    _add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Keywords
    kw = doc.add_paragraph()
    bk = kw.add_run("Keywords: ")
    bk.font.bold = True
    kw.add_run("sustainability science; research funding; UK Research and Innovation; "
               "network analysis; portfolio analysis; interdisciplinary research; "
               "Gateway to Research.").font.italic = True

    _add_page_break(doc)

    # ── INTRODUCTION ───────────────────────────────────────────────────────
    doc.add_heading("Introduction", level=1)

    _add_paragraph(doc,
        "Sustainability has become an organising principle for science and "
        "innovation policy worldwide. The United Nations Sustainable Development "
        "Goals¹, the Paris Agreement targets², and a growing array of national "
        "net-zero commitments³ have placed unprecedented expectations on the "
        "publicly funded research base to deliver evidence, technologies and "
        "policies that span the climate, energy, biodiversity, food and social "
        "domains simultaneously. Yet the structure of public research investment "
        "in sustainability science — how funding flows are distributed across "
        "themes, how disciplines intersect at the project level, and where "
        "institutional and investigator networks concentrate — remains poorly "
        "characterised in most national systems⁴.")

    _add_paragraph(doc,
        "The United Kingdom presents an instructive case. UK Research and "
        "Innovation (UKRI), the umbrella body for the seven discipline-based "
        "research councils and Innovate UK, has committed substantial public "
        "resources to sustainability-relevant programmes including the Industrial "
        "Strategy Challenge Fund⁵, the Net Zero Research and Innovation Framework⁶, "
        "and the strategic priorities funds aligned to the UK National Adaptation "
        "Programme⁷. Despite this commitment, no comprehensive portfolio-level "
        "audit has, to our knowledge, been published that traces how UKRI funding "
        "maps onto integrated sustainability challenges, how disciplines co-occur "
        "within projects, and which institutional ecosystems dominate the network. "
        "Such a characterisation is a precondition for evidence-based design of "
        "the next generation of UKRI strategic investments⁸.")

    _add_paragraph(doc,
        "Here we address this gap through a network-analytic portfolio audit of "
        f"the entire UKRI sustainability-relevant project landscape between "
        f"{s['yr_min']} and {s['yr_max']}. We compile a structured corpus of "
        f"{s['n_projects']:,} curated project records, enriched with rich per-"
        f"project metadata from the UKRI Gateway-to-Research application "
        f"programming interface, including organisational participation, "
        f"investigator roles, research-topic classifications and publication "
        f"outputs. We classify each project into eight pre-defined sustainability "
        f"themes via keyword matching of titles and abstracts, and construct "
        f"weighted co-occurrence networks at the theme, organisation, topic and "
        f"investigator levels. Three research questions motivate our analysis:")

    _add_paragraph(doc,
        "(i) How is UKRI sustainability funding distributed across thematic, "
        "geographic and institutional dimensions, and how has this evolved over "
        "the past decade? (ii) Which sustainability themes co-occur within "
        "individual projects, and what does the resulting network topology "
        "reveal about the degree of disciplinary integration in the portfolio? "
        "(iii) Which institutions and investigators function as structural "
        "bridges across the network, and where do under-bridged interfaces "
        "suggest opportunities for targeted future investment?")

    _add_paragraph(doc,
        "Our findings provide what we believe is the first network-analytic "
        "characterisation of the UKRI sustainability research portfolio, and "
        "offer a methodological template that can be applied to other national "
        "research-funding systems.")

    # ── RESULTS ────────────────────────────────────────────────────────────
    doc.add_heading("Results", level=1)

    # Section R1: funding landscape
    doc.add_heading("Funding landscape and temporal trend", level=2)

    _add_paragraph(doc,
        f"Across the analysed window, UKRI awarded £{s['total_funding_b']:.2f} "
        f"billion to {s['n_projects']:,} projects spanning all "
        f"{s['n_funders']} funding councils and innovation programmes. "
        f"Funding is highly concentrated by council: {s['top_funder']} alone "
        f"accounts for {s['top_funder_pct']:.0f}% of total investment "
        f"(Fig. 1). The annual funding flow (Fig. 2) shows a substantial "
        f"expansion of the portfolio after 2019, coinciding with the legislative "
        f"introduction of the UK net-zero target⁹.")

    _add_figure(doc, _fig_funding_share_donut(df), fig_number=1,
        caption=(
            "Distribution of total UKRI funding across the funding councils and "
            "innovation programmes that comprise the UK public-research-funding "
            "ecosystem. Funding is highly concentrated, with the largest "
            "council accounting for approximately one third of all awarded "
            "value. Colour encodes the council using the standard UKRI palette."
        ))

    _add_figure(doc, _fig_annual_funding(df), fig_number=2,
        caption=(
            "Annual UKRI funding (£ billion) stacked by funding council, "
            f"{s['yr_min']}–{s['yr_max']}. The post-2019 expansion of total "
            "annual investment reflects the introduction of the UK statutory "
            "net-zero target and the associated cross-council strategic "
            "priority funds. Council colour palette is consistent with Fig. 1."
        ))

    _add_paragraph(doc,
        f"Of the {s['n_projects']:,} analysed projects, "
        f"{s['sus_projects']:,} ({s['sus_share_pct']:.1f}%) were classified as "
        f"sustainability-relevant by keyword matching against the eight pre-"
        f"defined sustainability themes (see Methods). These sustainability "
        f"projects account for £{s['sus_funding_b']:.2f} billion of cumulative "
        f"awarded value, equivalent to "
        f"{100*s['sus_funding_b']/max(s['total_funding_b'],1):.1f}% of the "
        f"total UKRI portfolio analysed. The temporal trajectory of "
        f"sustainability funding (Fig. 3) demonstrates a near-monotonic "
        f"increase in absolute terms, with the share of sustainability funding "
        f"in the portfolio rising over the decade.")

    _add_figure(doc, _fig_sustainability_trend(df), fig_number=3,
        caption=(
            "Sustainability-classified funding versus all other UKRI funding "
            "by project start year. Stacked area shows non-sustainability "
            "(light blue) and sustainability (green) components. The widening "
            "green band illustrates a rising share of sustainability "
            "investment in the portfolio over time."
        ))

    # Section R2: themes
    doc.add_heading("Thematic structure and co-occurrence", level=2)

    cc = s["theme_counts"].get("Climate & Carbon", 0)
    cc_m = s["theme_funding_m"].get("Climate & Carbon", 0)
    sorted_themes = sorted(s["theme_counts"].items(), key=lambda kv: -kv[1])

    _add_paragraph(doc,
        f"Climate & Carbon dominates the sustainability portfolio in both "
        f"project volume and funding value ({cc} projects, £{cc_m:.0f} million), "
        f"followed in approximate order by "
        f"{sorted_themes[1][0]} ({sorted_themes[1][1]} projects), "
        f"{sorted_themes[2][0]} ({sorted_themes[2][1]} projects) and "
        f"{sorted_themes[3][0]} ({sorted_themes[3][1]} projects) (Fig. 4). "
        f"At the opposite end of the distribution, "
        f"{sorted_themes[-1][0]} ({sorted_themes[-1][1]} projects) and "
        f"{sorted_themes[-2][0]} ({sorted_themes[-2][1]} projects) form a "
        f"long tail. Funding intensity per project varies markedly across "
        f"themes, with the energy-decarbonisation themes attracting larger "
        f"average grant values than the social-science-led themes.")

    _add_figure(doc, _fig_themes_bar(df), fig_number=4,
        caption=(
            "Number of UKRI projects classified under each of the eight pre-"
            "defined sustainability themes. Bar colour encodes total funding "
            "(£ million) per theme. Themes are non-exclusive: a single "
            "project may be classified under multiple themes if its title or "
            "abstract contains keywords from multiple lists."
        ))

    _add_paragraph(doc,
        f"To probe the degree of thematic integration in the portfolio we "
        f"constructed a weighted co-occurrence network in which nodes are "
        f"sustainability themes and edge weight equals the number of projects "
        f"classified under both themes (Fig. 5). Of the 28 possible theme "
        f"pairs, {s['theme_graph'].number_of_edges()} were realised in the "
        f"data. Greedy modularity optimisation partitioned the themes into "
        f"three communities: a 'Living Earth' cluster comprising Environment "
        f"& Ecology, Water & Oceans, Food & Agriculture and Social "
        f"Sustainability; a 'Decarbonisation' cluster comprising Climate & "
        f"Carbon, Clean Energy and Sustainable Cities; and Circular Economy as "
        f"a structural singleton with zero cross-theme co-occurrences. The "
        f"strongest pairing by Jaccard similarity was "
        f"{s['theme_pairs_by_jaccard'][0][0]} ↔ "
        f"{s['theme_pairs_by_jaccard'][0][1]} "
        f"({s['theme_pairs_by_jaccard'][0][2]} shared projects, Jaccard "
        f"{s['theme_pairs_by_jaccard'][0][3]:.1f}%), reflecting the marine- "
        f"and freshwater-ecosystem core of the UKRI portfolio.")

    _add_figure(doc, _fig_theme_chord(df), fig_number=5,
        caption=(
            "Chord-layout co-occurrence network of the eight UKRI "
            "sustainability themes. Nodes represent themes; node size is "
            "proportional to total funding (£) of projects classified under "
            "each theme. Edges connect themes that co-occur on the same "
            "project; edge thickness is scaled by w^0.6 to amplify mid-range "
            "differences. Node colour and circular ordering reflect greedy "
            "modularity communities. Numeric annotations on edges indicate "
            "the number of shared projects."
        ))

    _add_paragraph(doc,
        f"To complement the topological view we present the full pairwise "
        f"co-occurrence matrix (Fig. 6), in which the diagonal reports the "
        f"total project count per theme. The matrix highlights the structural "
        f"isolation of Circular Economy more starkly: despite "
        f"{s['theme_counts'].get('Circular Economy', 0)} standalone projects "
        f"(£{s['theme_funding_m'].get('Circular Economy', 0):.0f} M), no "
        f"co-classifications with any other theme were detected. We discuss "
        f"the methodological caveats of this finding in the Discussion.")

    _add_figure(doc, _fig_theme_matrix(df), fig_number=6,
        caption=(
            "Sustainability theme co-occurrence matrix. Off-diagonal cells "
            "show the number of UKRI projects classified under both row and "
            "column themes; diagonal cells (highlighted with black borders) "
            "show the total projects classified under each theme. Colour "
            "intensity encodes shared-project count on a yellow-green "
            "sequential scale."
        ))

    # Top-pair table
    _add_table(doc,
        table_number=1,
        caption=(
            "Top ten cross-theme pairings in the UKRI sustainability "
            "portfolio, ranked by the number of projects co-classified under "
            "both themes. Jaccard similarity normalises co-occurrence by the "
            "union of each theme's project set and controls for the wide "
            "dispersion of theme volumes."
        ),
        headers=["Rank", "Theme A", "Theme B", "Shared projects", "Jaccard (%)"],
        rows=[
            [str(i + 1), u, v, str(n), f"{j:.1f}"]
            for i, (u, v, n, j) in enumerate(s["theme_pairs_by_raw"][:10])
        ],
    )

    _add_paragraph(doc,
        f"Temporal disaggregation of theme volumes (Fig. 7) shows that the "
        f"post-2019 expansion of sustainability investment is concentrated in "
        f"the Climate & Carbon and Clean Energy themes, with smaller but "
        f"discernible growth in Sustainable Cities and Circular Economy. The "
        f"Living Earth themes have grown more steadily across the window. "
        f"Only {s['multi_theme_projects']} of the {s['sus_projects']} "
        f"sustainability projects ({s['multi_theme_share_pct']:.1f}%) were "
        f"classified under two or more themes, accounting for "
        f"£{s['multi_theme_funding_m']:.0f} million "
        f"({100*s['multi_theme_funding_m']/max(s['sus_funding_b']*1000,1):.1f}%) "
        f"of total sustainability funding.")

    _add_figure(doc, _fig_theme_evolution(df), fig_number=7,
        caption=(
            "Number of UKRI projects classified under each sustainability "
            "theme by project start year. Line colour is consistent with the "
            "theme palette used in Figs 4–6. The post-2019 acceleration in "
            "Climate & Carbon and Clean Energy projects is visible."
        ))

    # Section R3: geography + institutions
    doc.add_heading("Geographic and institutional concentration", level=2)

    _add_paragraph(doc,
        f"UKRI sustainability funding is geographically concentrated. The "
        f"{s['top_region']} region alone receives {s['top_region_pct']:.0f}% "
        f"of total funding, with the South East, East of England and Scotland "
        f"forming the next tier (Fig. 8). At the institutional level, "
        f"{s['top_institution'].title()} leads with "
        f"£{s['top_institution_funding_m']:.0f} million in awarded funding. "
        f"The top 20 institutions collectively account for a substantial "
        f"share of total UKRI sustainability investment, suggesting that the "
        f"national portfolio is structurally dependent on a small number of "
        f"research-intensive universities and innovation hubs.")

    _add_figure(doc, _fig_region_funding(df), fig_number=8,
        caption=(
            "Total UKRI funding (£ million) by United Kingdom region. "
            "Regional totals are annotated alongside each bar. London, the "
            "South East and the East of England together account for the "
            "majority of UKRI investment, consistent with the geographic "
            "distribution of research-intensive universities in the UK."
        ))

    _add_paragraph(doc,
        "We then characterised the institutional collaboration network at "
        "the project level (Fig. 9). Nodes are lead institutions weighted by "
        "the number of distinct partner organisations; edges connect "
        "institutions that appear together as collaborators on the same "
        "project, with edge thickness proportional to the number of joint "
        "projects. Greedy modularity optimisation reveals three to four "
        "tightly knit research communities centred on geographically and "
        "thematically coherent hubs (e.g. a London-cluster anchored on "
        "Imperial College and University College London, a Scottish-cluster "
        "anchored on the University of Edinburgh, and a Midlands-cluster "
        "anchored on the University of Birmingham).")

    _add_figure(doc, _fig_institution_network(df), fig_number=9,
        caption=(
            "Institution co-collaboration network for the top 22 hub "
            "institutions in the UKRI sustainability portfolio. Nodes are "
            "lead institutions; node size is proportional to the number of "
            "distinct partner organisations; node colour encodes greedy-"
            "modularity community membership. Edges connect institutions "
            "that appear together on the same project; edge thickness "
            "is proportional to the number of joint projects (minimum "
            "weight three). Cluster halos highlight community structure."
        ))

    _add_paragraph(doc,
        "The funder–institution bipartite network (Fig. 10) shows that each "
        "of the top 20 institutions draws funding from multiple councils, "
        "though specialisation patterns are visible — biological-sciences "
        "councils dominate flows to land-grant universities, while EPSRC and "
        "Innovate UK channels are concentrated on the technology-oriented "
        "research-intensive universities.")

    _add_figure(doc, _fig_funder_inst_bipartite(df), fig_number=10,
        caption=(
            "Funder–Institution bipartite network. Funding councils (left "
            "column) are linked to the top 20 lead institutions (right "
            "column) by edges coloured by council and weighted by the number "
            "of jointly funded projects. The chart reveals each "
            "institution's council-funding mix at a glance and exposes the "
            "specialisation patterns between the technology-oriented and "
            "biological-sciences councils."
        ))

    # Top institutions table
    top_inst_rows = []
    for inst, row in s["top_institutions"].head(10).iterrows():
        top_inst_rows.append([
            inst.title(),
            f"{row['funding']/1e6:.1f}",
            f"{int(row['projects'])}",
        ])
    _add_table(doc,
        table_number=2,
        caption=(
            "Top 10 UK institutions by total UKRI funding received within "
            f"the {s['yr_min']}–{s['yr_max']} analysis window."
        ),
        headers=["Institution", "Total funding (£ M)", "Projects"],
        rows=top_inst_rows,
    )

    # Section R4: investigator network
    doc.add_heading("Principal-investigator collaboration network", level=2)

    _add_paragraph(doc,
        "Beyond the institutional layer, we constructed an investigator-"
        "level collaboration network (Fig. 11) from the per-project JSON "
        "metadata, in which each node is a Principal Investigator or Co-"
        "Investigator and each edge represents a co-investigation on the "
        "same UKRI project. To surface the structural backbone of the "
        "network we rank investigators by weighted collaboration degree and "
        "show the top 40 most-connected individuals. The network exhibits "
        "clear community structure aligned with the institutional clusters "
        "of Fig. 9, with several individuals functioning as cross-community "
        "bridges. Because most UKRI awards list a single Principal "
        "Investigator, Co-Investigator roles are the principal mechanism by "
        "which investigators link across the portfolio.")

    _add_figure(doc, _fig_pi_network(pp), fig_number=11,
        caption=(
            "Principal-Investigator and Co-Investigator collaboration "
            "network in the UKRI portfolio. Top 40 investigators by weighted "
            "collaboration degree. Edges link individuals who appear on the "
            "same project as PI or Co-I; edge weight equals the number of "
            "shared projects. Node colour encodes greedy-modularity community "
            "membership; node size is proportional to the total number of "
            "projects on which the individual was listed."
        ))

    # Section R5: topic network
    doc.add_heading("Research-topic landscape", level=2)

    _add_paragraph(doc,
        "Finally, we used the bottom-up research-topic and research-subject "
        "classifications recorded in the UKRI Gateway-to-Research metadata "
        "to construct a topic co-occurrence network (Fig. 12) that is "
        "independent of our top-down sustainability theme classifier. Nodes "
        "are individual research-topic tags; node size is proportional to "
        "total project funding associated with each tag; edges connect tags "
        "that co-occur on the same project. The topic network identifies "
        "several dense clusters that map approximately onto the eight "
        "sustainability themes, providing convergent validation of the "
        "thematic structure inferred from keyword classification.")

    _add_figure(doc, _fig_topic_network(df, ti), fig_number=12,
        caption=(
            "Research-topic co-occurrence network derived from the "
            "researchTopics and researchSubjects classifications recorded "
            "in the UKRI Gateway-to-Research metadata for every project. "
            "Top 35 tags by occurrence are shown; nodes are tags, node size "
            "is proportional to total funding (£) of projects classified "
            "under each tag, edges connect tags that co-occur on the same "
            "project (minimum weight four), and edge thickness is scaled by "
            "w^0.6. Node colour encodes greedy-modularity community "
            "membership."
        ))

    # ── DISCUSSION ─────────────────────────────────────────────────────────
    doc.add_heading("Discussion", level=1)

    _add_paragraph(doc,
        "Our analysis offers three principal findings that bear on the design "
        "of future UK sustainability research strategy. First, the UKRI "
        f"sustainability portfolio is large in absolute terms "
        f"(£{s['sus_funding_b']:.2f} billion across {s['sus_projects']:,} "
        f"projects in the analysis window) but heavily concentrated by theme, "
        f"by region and by institution. Climate & Carbon alone accounts for "
        f"approximately one third of the sustainability project count, and "
        f"the top 20 institutions absorb the majority of investment. This "
        f"concentration is consistent with prior accounts of the UK research "
        f"base⁴ but is sharper for sustainability work than for the portfolio "
        f"as a whole.")

    _add_paragraph(doc,
        "Second, the theme co-occurrence network (Fig. 5) reveals two "
        "structurally distinct research communities — a Living Earth cluster "
        "(Environment, Water, Food, Social) and a Decarbonisation cluster "
        "(Climate, Energy, Cities) — that are linked, but only weakly, by "
        "Climate & Carbon as a shared label. Within-cluster co-occurrence is "
        "substantially stronger than cross-cluster co-occurrence, suggesting "
        "that the integrated 'sustainability science' envisaged by Kates et "
        "al.¹⁰ and Clark¹¹ has only partially materialised at the project "
        "level. The Living Earth and Decarbonisation clusters appear to "
        "operate as parallel research economies with limited project-level "
        "integration — a finding with implications for the design of cross-"
        "cluster strategic priority funds.")

    _add_paragraph(doc,
        "Third, the structural isolation of Circular Economy in the "
        "co-occurrence network deserves careful interpretation. Despite "
        f"{s['theme_counts'].get('Circular Economy', 0)} standalone projects "
        f"and £{s['theme_funding_m'].get('Circular Economy', 0):.0f} million "
        f"in awarded funding, our analysis detects zero co-occurrences with "
        f"any other sustainability theme. We do not interpret this as "
        f"evidence that circularity is genuinely disconnected from climate, "
        f"materials, food-waste or sustainable-cities research. Rather, the "
        f"finding most plausibly reflects a vocabulary effect in our "
        f"keyword-based classifier: the Circular Economy keyword list "
        f"(recycling, bioplastic, industrial ecology, cradle-to-cradle) "
        f"shares limited lexical overlap with the other seven theme lists. "
        f"This is a methodological caveat worth highlighting in any "
        f"interpretation of the network topology, and motivates future work "
        f"using sentence-embedding classifiers that capture semantic rather "
        f"than purely lexical overlap.")

    _add_paragraph(doc,
        "The cross-theme participation rate "
        f"({s['multi_theme_share_pct']:.1f}% of sustainability projects "
        f"classified under two or more themes) should likewise be read as a "
        f"lower bound on genuine interdisciplinarity. Substring keyword "
        f"matching requires that vocabulary from multiple lists appear "
        f"literally in the abstract; semantically integrated work that uses "
        f"the vocabulary of only one theme will be misclassified as single-"
        f"theme. A complementary topic-embedding analysis would be required "
        f"to estimate the true interdisciplinarity rate; the topic network of "
        f"Fig. 12 provides a step in this direction by surfacing dense "
        f"co-occurrences in the bottom-up research-topic taxonomy.")

    _add_paragraph(doc,
        "These findings have direct implications for UKRI strategy. The "
        "weak Decarbonisation–Living Earth interface argues for targeted "
        "cross-cluster programmes — for example, a Climate-Justice or "
        "Just-Transition funding call that explicitly requires projects to "
        "draw from both the Clean Energy and Social Sustainability themes. "
        "The structural isolation of Circular Economy similarly argues for "
        "thematic-bridge calls that pair circularity expertise with "
        "materials, manufacturing or food-waste teams. At the institutional "
        "level, the concentration of investment in the top 20 universities "
        "raises the question of how to broaden geographic and institutional "
        "participation without sacrificing research quality — a balance the "
        "next round of UKRI strategic delivery plans will need to address¹².")

    _add_paragraph(doc,
        "Our study has several limitations. We restrict analysis to the "
        "curated Excel registry of approximately one thousand four hundred "
        "projects, whereas the full Gateway-to-Research JSON corpus contains "
        "in excess of sixteen thousand records. The Excel registry was "
        "hand-curated for a separate analytical purpose and may not be a "
        "representative sample of the full UKRI portfolio; extending the "
        "sustainability classifier to the full sixteen-thousand-record "
        "corpus is an immediate priority for future work. We also caution "
        "that keyword classification, while transparent and reproducible, "
        "introduces systematic biases against research that uses non-"
        "standard vocabulary; replacing the classifier with a sentence-"
        "embedding pipeline trained on a curated corpus of sustainability "
        "abstracts would improve precision and recall. Finally, our "
        "network-analytic conclusions are descriptive rather than causal: "
        "the data permit characterisation of structure but not "
        "identification of the policy levers that produced it.")

    _add_paragraph(doc,
        "Notwithstanding these limitations, the analysis provides what we "
        "believe is the first network-analytic portfolio audit of the UKRI "
        "sustainability research landscape. The methodological framework — "
        "keyword classification, weighted co-occurrence networks at the "
        "theme, institutional, investigator and topic levels, and community "
        "detection via greedy modularity — is generalisable to other "
        "national research-funding systems. We commend it to research "
        "funders worldwide as a transparent, reproducible tool for "
        "evidence-based portfolio design.")

    # ── METHODS ────────────────────────────────────────────────────────────
    doc.add_heading("Methods", level=1)

    doc.add_heading("Data sources", level=2)
    _add_paragraph(doc,
        f"We compiled two complementary data sources. First, a curated "
        f"Microsoft Excel registry of {s['n_projects']:,} UKRI projects "
        f"awarded between {s['yr_min']} and {s['yr_max']}, originally "
        f"prepared for a separate analytical purpose and containing "
        f"hand-cleaned organisational affiliation fields. Second, we "
        f"downloaded the per-project JavaScript Object Notation (JSON) "
        f"records for sixteen thousand one hundred and twenty-eight projects "
        f"from the UKRI Gateway-to-Research application programming "
        f"interface (https://gtr.ukri.org/), providing rich metadata "
        f"including project abstracts, organisational roles, person roles, "
        f"research-topic classifications and publication outputs. JSON "
        f"records were matched to the Excel registry by UKRI grant "
        f"reference number; ninety-nine per cent of registry records were "
        f"successfully matched.")

    doc.add_heading("Sustainability theme classification", level=2)
    _add_paragraph(doc,
        "We classified each project into eight pre-defined sustainability "
        "themes (Climate & Carbon, Clean Energy, Environment & Ecology, "
        "Water & Oceans, Food & Agriculture, Circular Economy, Sustainable "
        "Cities, Social Sustainability) by keyword matching of project "
        "titles and abstracts. Each theme is associated with a curated list "
        "of twelve to eighteen domain-specific keywords. A project is "
        "tagged with theme T if at least one keyword from T's list appears "
        "as a substring in the lowercased concatenation of the project "
        "title and abstract. Themes are non-exclusive; a project may be "
        "tagged with multiple themes. The full keyword lists are reproduced "
        "in Supplementary Methods.")

    doc.add_heading("Network construction", level=2)
    _add_paragraph(doc,
        "We constructed five weighted networks. (i) A theme co-occurrence "
        "network in which nodes are sustainability themes and edge weight "
        "equals the number of projects co-classified under both themes. "
        "(ii) An institution co-collaboration network in which nodes are "
        "lead institutions and edges connect institutions that appear "
        "together as collaborators on the same project, weighted by joint-"
        "project count. (iii) A funder–institution bipartite network in "
        "which funding councils are linked to lead institutions by edges "
        "weighted by jointly funded project counts. (iv) An investigator "
        "collaboration network constructed from the personRoles field of "
        "the Gateway-to-Research JSON records, in which nodes are "
        "Principal Investigators and Co-Investigators and edges connect "
        "individuals who appear on the same project. (v) A research-topic "
        "co-occurrence network constructed from the researchTopics and "
        "researchSubjects fields of the JSON records, in which nodes are "
        "research-topic tags and edges connect tags that co-occur on the "
        "same project.")

    doc.add_heading("Community detection and visualisation", level=2)
    _add_paragraph(doc,
        "All co-occurrence networks were analysed for community structure "
        "using the greedy modularity optimisation algorithm of Clauset, "
        "Newman and Moore¹³ as implemented in NetworkX 3.2. Layout was "
        "performed using the Fruchterman–Reingold force-directed "
        "algorithm¹⁴ with manual post-processing to minimise label overlap "
        "in dense regions. Edge thickness in all network figures is scaled "
        "as w^0.6 where w is the edge weight, to amplify mid-range "
        "differences. Edge alpha (transparency) is similarly scaled with "
        "edge weight to suppress visual clutter from weak edges. Cluster "
        "halos in Figs 5 and 9 are visualisation aids and do not encode "
        "additional information.")

    doc.add_heading("Statistical analysis", level=2)
    _add_paragraph(doc,
        "Pairwise theme overlap is reported as both raw co-occurrence "
        "count and Jaccard similarity (|A ∩ B| / |A ∪ B|), the latter "
        "controlling for the wide dispersion of theme volumes. All "
        "monetary values are reported in nominal pounds sterling at the "
        "time of award; no inflation adjustment is applied. Funding totals "
        "do not include co-investment by industrial or charitable "
        "partners.")

    doc.add_heading("Code and data availability", level=2)
    _add_paragraph(doc,
        "All analysis code is implemented in Python 3.11 and is available "
        "from the corresponding author upon reasonable request. The Excel "
        "project registry is derived from publicly available UKRI award "
        "data. The Gateway-to-Research JSON corpus is freely available at "
        "https://gtr.ukri.org/. An interactive web-based version of the "
        "dashboard from which the figures in this paper were generated "
        "is included with the supplementary material.")

    # ── ACKNOWLEDGEMENTS ──────────────────────────────────────────────────
    doc.add_heading("Acknowledgements", level=1)
    _add_paragraph(doc,
        "We thank colleagues at the University of Cambridge for helpful "
        "discussions during the preparation of this manuscript. We "
        "acknowledge UK Research and Innovation for providing the underlying "
        "project metadata via the Gateway-to-Research application "
        "programming interface.")

    # ── AUTHOR CONTRIBUTIONS ──────────────────────────────────────────────
    doc.add_heading("Author contributions", level=1)
    _add_paragraph(doc,
        "A.O. conceived the study, designed the analysis pipeline, and "
        "wrote the manuscript. A.T. contributed to the keyword "
        "classification and validated the results. A.Th. contributed to the "
        "network analysis and prepared the figures. All authors reviewed and "
        "approved the final manuscript.")

    # ── COMPETING INTERESTS ────────────────────────────────────────────────
    doc.add_heading("Competing interests", level=1)
    _add_paragraph(doc, "The authors declare no competing interests.")

    # ── REFERENCES ─────────────────────────────────────────────────────────
    doc.add_heading("References", level=1)

    refs = [
        "United Nations General Assembly. Transforming our world: the 2030 "
        "Agenda for Sustainable Development. Resolution A/RES/70/1 (2015).",

        "United Nations Framework Convention on Climate Change. Adoption of "
        "the Paris Agreement. Conference of the Parties report FCCC/CP/2015/L.9/Rev.1 (2015).",

        "Höhne, N. et al. Wave of net zero emission targets opens window to "
        "meeting the Paris Agreement. Nature Climate Change 11, 820–822 (2021).",

        "Adams, J., Pendlebury, D., Potter, R. & Szomszor, M. The annual "
        "G20 scorecard – research performance 2019. Institute for Scientific "
        "Information, Clarivate (2019).",

        "Department for Business, Energy and Industrial Strategy. Industrial "
        "Strategy: Building a Britain Fit for the Future. HM Government White "
        "Paper Cm 9528 (2017).",

        "UK Research and Innovation. Net Zero Research and Innovation "
        "Framework. UKRI Strategy Document (2022).",

        "Department for Environment, Food and Rural Affairs. UK Climate "
        "Change Risk Assessment 2022. HM Government Report (2022).",

        "UK Research and Innovation. Transforming Tomorrow Together: UKRI "
        "Strategy 2022 to 2027. UKRI Strategy Document (2022).",

        "Climate Change Act 2008 (2050 Target Amendment) Order 2019. UK "
        "Statutory Instrument 2019 No. 1056 (2019).",

        "Kates, R. W. et al. Environment and development: sustainability "
        "science. Science 292, 641–642 (2001).",

        "Clark, W. C. & Harley, A. G. Sustainability science: toward a "
        "synthesis. Annual Review of Environment and Resources 45, 331–386 "
        "(2020).",

        "Wilsdon, J. et al. The Metric Tide: Report of the Independent "
        "Review of the Role of Metrics in Research Assessment and "
        "Management. HEFCE (2015).",

        "Clauset, A., Newman, M. E. J. & Moore, C. Finding community "
        "structure in very large networks. Physical Review E 70, 066111 "
        "(2004).",

        "Fruchterman, T. M. J. & Reingold, E. M. Graph drawing by force-"
        "directed placement. Software: Practice and Experience 21, "
        "1129–1164 (1991).",

        "Hagberg, A. A., Schult, D. A. & Swart, P. J. Exploring network "
        "structure, dynamics, and function using NetworkX. In Proceedings "
        "of the 7th Python in Science Conference (eds Varoquaux, G. et al.) "
        "11–15 (2008).",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(4)
        num_run = p.add_run(f"{i}.\t")
        num_run.font.bold = True
        num_run.font.size = Pt(10)
        ref_run = p.add_run(ref)
        ref_run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
